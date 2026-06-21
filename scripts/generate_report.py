import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_report():
    print("Initializing professional expanded document generator...")
    doc = docx.Document()
    
    # 1. Margins Setup (1 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Add page numbering/headers
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "HITEC Smart University Portal (HiSUP) — ADMS Project Report"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.style.font.name = 'Calibri'
        hp.style.font.size = Pt(9.0)
        hp.style.font.italic = True
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = "Department of Computer Science, HITEC University Taxila"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.style.font.name = 'Calibri'
        fp.style.font.size = Pt(9.0)

    # 2. Modify Styles for a Professional Look
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(44, 62, 80)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    
    # Colors
    PRIMARY_COLOR = RGBColor(26, 82, 118)    # Deep Steel Blue
    SECONDARY_COLOR = RGBColor(40, 116, 166)  # Muted Blue
    
    def add_styled_heading(text, level, space_before=12, space_after=6):
        heading = doc.add_heading(text, level=level)
        h_style = heading.style
        h_font = h_style.font
        h_font.name = 'Arial'
        h_font.bold = True
        
        if level == 1:
            h_font.size = Pt(18)
            h_font.color.rgb = PRIMARY_COLOR
            heading.paragraph_format.space_before = Pt(space_before + 6)
            heading.paragraph_format.space_after = Pt(space_after + 4)
            heading.paragraph_format.keep_with_next = True
        elif level == 2:
            h_font.size = Pt(14)
            h_font.color.rgb = SECONDARY_COLOR
            heading.paragraph_format.space_before = Pt(space_before)
            heading.paragraph_format.space_after = Pt(space_after + 2)
            heading.paragraph_format.keep_with_next = True
        return heading

    def add_p(text, bold_prefix=None, italic=False, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.bold = True
            run_b.font.color.rgb = RGBColor(44, 62, 80)
            
        run_t = p.add_run(text)
        run_t.font.italic = italic
        run_t.font.color.rgb = RGBColor(44, 62, 80)
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.bold = True
            
        p.add_run(text)
        return p

    def add_code_block(code_text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        
        # Shade cell background (#F4F6F7)
        shading_elm = parse_xml(r'<w:shd {} w:fill="F4F6F7"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Left border (#A6ACAF, size 24 = 3pt)
        tcPr = cell._tc.get_or_add_tcPr()
        borders_elm = parse_xml(r'''<w:tcBorders {}>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="A6ACAF"/>
            <w:top w:val="none"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>'''.format(nsdecls('w')))
        tcPr.append(borders_elm)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(33, 47, 60)
        
        # Add space after table
        p_after = doc.add_paragraph()
        p_after.paragraph_format.space_before = Pt(0)
        p_after.paragraph_format.space_after = Pt(4)
        p_after.paragraph_format.line_spacing = Pt(2)

    def add_screenshot_placeholder(label):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        
        # Shade cell background
        shading_elm = parse_xml(r'<w:shd {} w:fill="EAECEE"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Outer borders
        tcPr = cell._tc.get_or_add_tcPr()
        borders_elm = parse_xml(r'''<w:tcBorders {}>
            <w:left w:val="single" w:sz="12" w:space="0" w:color="BDC3C7"/>
            <w:top w:val="single" w:sz="12" w:space="0" w:color="BDC3C7"/>
            <w:bottom w:val="single" w:sz="12" w:space="0" w:color="BDC3C7"/>
            <w:right w:val="single" w:sz="12" w:space="0" w:color="BDC3C7"/>
        </w:tcBorders>'''.format(nsdecls('w')))
        tcPr.append(borders_elm)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(32)
        p.paragraph_format.space_after = Pt(32)
        
        run = p.add_run(f"📷 [PLACEHOLDER SCREENSHOT: {label}]")
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(127, 140, 141)
        
        p_caption = doc.add_paragraph()
        p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_caption.paragraph_format.space_before = Pt(4)
        p_caption.paragraph_format.space_after = Pt(12)
        run_cap = p_caption.add_run(f"Figure: Demonstration of {label}")
        run_cap.font.size = Pt(9.5)
        run_cap.font.italic = True
        run_cap.font.color.rgb = RGBColor(127, 140, 141)

    # PAGE 1: TITLE PAGE
    print("Writing Title Page...")
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(96)
    title_p.paragraph_format.space_after = Pt(18)
    title_run = title_p.add_run("HITEC SMART UNIVERSITY PORTAL (HiSUP)")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(26)
    title_run.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(72)
    sub_run = subtitle_p.add_run("Advanced Database Management Systems (CS-402) — Semester Project Report")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = SECONDARY_COLOR

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_after = Pt(96)
    
    info_run1 = info_p.add_run("Prepared by:\n")
    info_run1.bold = True
    info_run2 = info_p.add_run("Student Name: Ibrahim / Zubair\nRoll Number: [Your Roll Number]\n\n")
    info_run3 = info_p.add_run("Submitted to:\n")
    info_run3.bold = True
    info_run4 = info_p.add_run("Department of Computer Science\nHITEC University, Taxila\n\nDate: Spring Semester 2026")
    
    for r in [info_run1, info_run2, info_run3, info_run4]:
        r.font.name = 'Calibri'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(44, 62, 80)
        
    links_p = doc.add_paragraph()
    links_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links_run = links_p.add_run("GitHub Repository: [INSERT GITHUB URL HERE]\nLive Application URL: [INSERT LIVE DEPLOYMENT URL HERE]")
    links_run.font.name = 'Consolas'
    links_run.font.size = Pt(10)
    links_run.font.color.rgb = RGBColor(40, 116, 166)
    
    doc.add_page_break()

    # PAGE 2: TABLE OF CONTENTS
    print("Writing Table of Contents placeholder...")
    add_styled_heading("Table of Contents", level=1)
    add_p("This section contains the project report Table of Contents. Open this document in Microsoft Word, right-click inside the field placeholder below, and choose 'Update Field' to dynamically construct the headings list with their corresponding page numbers.", italic=True)
    
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.space_before = Pt(24)
    toc_p.paragraph_format.space_after = Pt(24)
    toc_run = toc_p.add_run("[DYNAMIC TABLE OF CONTENTS FIELD - UPDATE IN MS WORD]")
    toc_run.bold = True
    toc_run.font.color.rgb = RGBColor(127, 140, 141)
    
    doc.add_page_break()

    # PAGE 3-5: EXECUTIVE SUMMARY & SYSTEM ROLES
    print("Writing Chapter 1: Executive Summary & Roles...")
    add_styled_heading("1. Executive Summary & System Roles", level=1)
    add_p("The HITEC Smart University Portal (HiSUP) is an enterprise-class academic enterprise resource planning (ERP) platform. Designed for modern academic operations, HiSUP handles core university operational workflows, including detailed student records, multi-level academic program catalogs, course registration with dynamic seat checks, grade processing, student attendance monitoring, financial fee payments, university library management, hostel assignments, exam schedules, and extensive audit tracking.", bold_prefix="System Purpose: ")
    add_p("In terms of software architecture, while the original course description assumed an ASP.NET Core 8 with SQL Server stack, this project is built using a highly efficient Node.js/Express backend, a modular React frontend styled with Bootstrap 5, and an enterprise Oracle 12c+ relational database. The interface is responsive and provides dedicated interfaces for four major roles: Administrators, Faculty Members, Finance Cashiers, and Students. Business logic is strictly managed at the database tier using stored procedures, database functions, triggers, and Virtual Private Database (VPD) row-level security policy constructs. Communication between the Express server and Oracle is achieved through the binary node-oracledb client driver.", bold_prefix="Technology Stack & Architecture: ")
    
    add_styled_heading("1.1 User Roles & Functionalities", level=2)
    add_p("Access control and feature visibility in HiSUP are divided into four main roles, each corresponding to a distinct operational profile in the university structure:", bold_prefix="Role Division: ")
    
    add_p("Administrators act as the system configuration operators. They configure core department structures, program curricula, catalog courses, establish academic sections, create user login credentials, assign faculty members to sections, and monitor administrative logs. They have unrestricted read/write permission to audit logs and base master records. This allows the administrative staff to manage the entire structure of the portal from a central control panel.", bold_prefix="1. Admin Role: ")
    add_screenshot_placeholder("Admin Dashboard Portal")
    
    add_p("Students can access a comprehensive dashboard showing their cumulative GPA (CGPA), class attendance summaries, and outstanding fee balances. The student interface allows course registration (with dynamic, lock-protected seat validation), viewing fee structures, processing fee payments semester-wise, checking out library books, returning library items, and viewing their courses and grades grouped by previous semesters along with their Semester GPA (SGPA) and overall CGPA. Downloading transcripts has been disabled for security/policy reasons. It is designed to be user-friendly and fully self-service.", bold_prefix="2. Student Role: ")
    add_screenshot_placeholder("Student Dashboard Portal")
    
    add_p("Faculty members utilize the portal to manage academic classes assigned to them. They can view student rosters, record daily attendance for up to 16 lectures per term, input exam grades, update student score cards, and view workload ranking reports that detail their section capacities and enrolled student counts. This facilitates real-time entry of performance indicators.", bold_prefix="3. Faculty Role: ")
    add_screenshot_placeholder("Faculty Dashboard Portal")
    
    add_p("Finance cashiers handle university financial operations. They configure semester tuition structures, post fee invoices, process student payments, issue receipt vouchers, and review outstanding accounts using fee defaulter views to trace students who have unpaid balances. This module integrates transaction controls to prevent concurrency conflicts.", bold_prefix="4. Finance Role: ")
    add_screenshot_placeholder("Finance Defaulters Interface")

    doc.add_page_break()

    # PAGE 6-9: RELATIONAL DATABASE TABLES
    print("Writing Chapter 2: Relational Database Tables...")
    add_styled_heading("2. Relational Database Tables", level=1)
    add_p("The database schema for HiSUP is designed to support the academic operational demands of a high-enrollment institution. In total, 20 interrelated tables are defined. Relationships are enforced via primary keys and foreign key constraints with strict ON DELETE CASCADE and ON DELETE SET NULL behaviors to prevent orphans and ensure referential integrity. In addition, CHECK constraints, DEFAULT constraints, and UNIQUE requirements enforce domain integrity across all attributes. Below is a detailed review of each table's schema, constraints, and operational role in the database.", bold_prefix="Relational Overview: ")
    
    tables_list = [
        ("departments", "Stores department master records including unique names and codes (e.g. Computer Science, CS). It is referenced by courses and programs to maintain department groupings."),
        ("programs", "Stores degree programs (e.g. BS Computer Science) linked to specific parent departments. Includes duration in semesters and references departments with ON DELETE CASCADE."),
        ("user_accounts", "Acts as the central credentials directory, logging email, hashed/plain passwords, and roles (Admin, Student, Finance, Faculty) for authentication checks."),
        ("students", "Stores student profiles, program links, and encrypted CNIC identifiers for personal security. Links directly to user_accounts via user_id."),
        ("faculty", "Stores faculty credentials, department associations, and encrypted bank account strings for routing payroll. Links to user_accounts."),
        ("staff", "Stores profiles and department assignments for administrative staff members who have access to backend records."),
        ("courses", "Maintains the course catalog, credit hour assignments, and self-referential prerequisite course chains for curriculum mapping."),
        ("sections", "Tracks active classes, referencing courses, faculty, term, capacity, room, and available seats. Enforces seat count limits via CHECK constraints."),
        ("enrollments", "Links students to course sections, enforcing a unique constraint on student-section pairs to prevent double registration."),
        ("grades", "Maintains academic grade values, letter scores, and GPA points earned for each enrollment. Linked to enrollments with ON DELETE CASCADE."),
        ("attendance_records", "Tracks daily attendance statuses (Present, Absent, Late) per enrollment for lectures 1 to 16. Enforces unique lecture numbers."),
        ("fee_structures", "Establishes tuition fee amounts due per program and semester, including due dates, referenced by programs."),
        ("fee_payments", "Logs payments processed for students, payment methods, bank references, and encrypted bank credentials. Updates student balances."),
        ("library_items", "Tracks inventory details for library books and journals, including total and available copies. Used for borrowing validations."),
        ("library_issues", "Logs book checkout events, due dates, return dates, checkout statuses, and calculated overdue fines. Links students and library items."),
        ("hostels", "Tracks residential hostel halls and total/available dorm rooms to manage on-campus living accommodations."),
        ("hostel_allotments", "Links resident students with specific hostel rooms and allotment start/end dates. Prevents room over-allocation."),
        ("exam_schedule", "Tracks final and midterm exam dates, times, rooms, and section associations for student scheduling."),
        ("results", "Acts as the static term-GPA log for historical academic reporting and transcript extraction."),
        ("audit_log", "A system-wide audit repository logging table name, old/new values, changed_by, and timestamp details via system triggers.")
    ]
    
    for name, desc in tables_list:
        add_styled_heading(f"2.{tables_list.index((name, desc))+1} Table: {name}", level=2)
        add_p(desc, bold_prefix="Operational Role: ")
        
    doc.add_page_break()

    # PAGE 10-18: DATABASE PROGRAMMABILITY OBJECTS
    print("Writing Chapter 3: Database Objects...")
    add_styled_heading("3. Database Objects Overview", level=1)
    add_p("Business logic in HiSUP is strictly maintained at the database tier using stored procedures, database functions, triggers, and views. This centralizes validation rules, optimizes execution plans, and secures data tables against raw queries. Below is a detailed breakdown of each programmability object, including source code excerpts for the most critical modules.", bold_prefix="Programmability Architecture: ")
    
    # Procedures
    add_styled_heading("3.1 Stored Procedures", level=2)
    add_p("Stored procedures encapsulate complex transactional workflows directly on the database server. This reduces network round-trips, optimizes execution, and secures access. The schema defines exactly 16 stored procedures:", bold_prefix="Purpose: ")
    
    procedures_list = [
        ("RegisterStudent", "Registers a student by inserting credentials into user_accounts and generating a student demographic record inside a transaction."),
        ("EnrollInCourse", "Registers a student in a section, validating seat availability and locking the section row for concurrency control."),
        ("ProcessFeePayment", "Processes student payment, inserts a fee_payments record, and commits the transaction."),
        ("AuthenticateUser", "Verifies email and password credentials, returning the user's role and reference ID."),
        ("GenerateTranscriptPDF", "Queries course grades and CGPA and returns a text-based transcript summary buffer."),
        ("CalculateSemesterGPA", "Computes the term GPA by averaging grade points earned in sections for a given semester."),
        ("MarkAttendance", "Upserts class attendance statuses for a section using input JSON payloads."),
        ("AllocateHostelRoom", "Allots a hostel room to a student, reducing the available room count in that hostel."),
        ("IssueLibraryBook", "Checks out a book, validating student status and decrementing the available copy count."),
        ("ReturnLibraryBook", "Records book return, computes overdue fines if applicable, and increments available copies."),
        ("AddExamResult", "Upserts student final grades into the grades table using a MERGE statement."),
        ("GetStudentReport", "Queries academic course records, grades, and GPAs for student reporting."),
        ("GetFacultyWorkload", "Compiles teaching workloads and sections assigned to a faculty member, ranking them by size."),
        ("GetDepartmentEnrollment", "Calculates total enrollment counts and unique student registrations in a department."),
        ("GenerateFeeSlip", "Queries fee structures and payment logs to print a detailed student fee invoice."),
        ("SearchCourses", "Dynamically searches catalog courses matching optional department and keyword filters safely.")
    ]
    for name, desc in procedures_list:
        add_p(desc, bold_prefix=f"* Procedure: `{name}` — ")
        
    add_styled_heading("3.1.1 Stored Procedure Source Code Examples", level=2)
    add_p("Below is the source code for the core transactional stored procedures, demonstrating Oracle PL/SQL parameters, row locking, and transaction controls:", bold_prefix="Implementation: ")
    
    add_p("RegisterStudent Stored Procedure:", bold_prefix="1. ")
    add_code_block("""CREATE OR REPLACE PROCEDURE RegisterStudent(
  p_first_name IN VARCHAR2,
  p_last_name IN VARCHAR2,
  p_email IN VARCHAR2,
  p_cnic IN VARCHAR2,
  p_program_id IN NUMBER,
  p_out_student_id OUT NUMBER
) AS
BEGIN
  IF p_first_name IS NULL OR p_last_name IS NULL OR p_email IS NULL OR p_program_id IS NULL THEN
    RAISE_APPLICATION_ERROR(-20001, 'First name, last name, email, and program are required.');
  END IF;

  INSERT INTO user_accounts(email, password_hash, role, reference_id)
  VALUES (p_email, 'PLACEHOLDER_HASH', 'Student', NULL)
  RETURNING user_id INTO p_out_student_id;

  INSERT INTO students(user_id, first_name, last_name, email, cnic_raw, cnic_encrypted, program_id)
  VALUES (p_out_student_id, p_first_name, p_last_name, p_email, p_cnic, fn_encrypt_value(p_cnic), p_program_id)
  RETURNING student_id INTO p_out_student_id;
EXCEPTION
  WHEN DUP_VAL_ON_INDEX THEN
    RAISE_APPLICATION_ERROR(-20002, 'A student or email already exists.');
  WHEN OTHERS THEN
    ROLLBACK;
    RAISE;
END;""")

    add_p("EnrollInCourse Stored Procedure:", bold_prefix="2. ")
    add_code_block("""CREATE OR REPLACE PROCEDURE EnrollInCourse(
  p_student_id IN NUMBER,
  p_section_id IN NUMBER
) AS
  l_available NUMBER;
  l_enrollments NUMBER;
BEGIN
  IF p_student_id IS NULL OR p_section_id IS NULL THEN
    RAISE_APPLICATION_ERROR(-20001, 'Student ID and section ID are required.');
  END IF;

  -- Concurrency check: lock section row during seat checks
  SELECT available_seats INTO l_available FROM sections WHERE section_id = p_section_id FOR UPDATE;
  IF l_available <= 0 THEN
    RAISE_APPLICATION_ERROR(-20003, 'No seats available in this section.');
  END IF;

  SELECT COUNT(*) INTO l_enrollments FROM enrollments WHERE student_id = p_student_id AND section_id = p_section_id;
  IF l_enrollments > 0 THEN
    RAISE_APPLICATION_ERROR(-20004, 'Student already enrolled in this section.');
  END IF;

  INSERT INTO enrollments (student_id, section_id) VALUES (p_student_id, p_section_id);
  UPDATE sections SET available_seats = available_seats - 1 WHERE section_id = p_section_id;

  COMMIT;
EXCEPTION
  WHEN OTHERS THEN
    ROLLBACK;
    IF SQLCODE = -20003 OR SQLCODE = -20004 THEN
      RAISE;
    ELSE
      RAISE_APPLICATION_ERROR(-20005, 'Enrollment failed: ' || SQLERRM);
    END IF;
END;""")
        
    # Functions
    add_styled_heading("3.2 User-Defined Functions", level=2)
    add_p("User-defined functions (UDFs) return values and allow calculated values to be embedded directly inside SQL SELECT queries. The system defines 8 functions:", bold_prefix="Purpose: ")
    functions_list = [
        ("fn_encrypt_value", "Encrypts raw inputs using Base64 encoding to secure sensitive database attributes."),
        ("fn_decrypt_value", "Decrypts Base64 encrypted strings back into readable plain text parameters."),
        ("fn_GetLetterGrade", "Translates numeric student scores into official letter grades (e.g. A+, B, F)."),
        ("fn_CalculateGradePoints", "Translates numeric student scores into grade points (e.g. 4.0, 3.0, 0.0) for GPA calculations."),
        ("fn_GetOutstandingFee", "Computes outstanding student balance by subtracting payments from program fee structures."),
        ("fn_GetSemesterOutstanding", "Computes the outstanding fee balance for a student for a specific semester by subtracting approved semester payments from program fee structures."),
        ("fn_GetAttendancePercentage", "Calculates the student's attendance percentage based on present vs absent lecture counts."),
        ("fn_IsLibraryItemAvailable", "Checks copy inventories to return whether a library book is available for borrowing."),
        ("fn_CalculateCGPA", "Calculates the student's cumulative CGPA weighted by course credit hours across all semesters.")
    ]
    for name, desc in functions_list:
        add_p(desc, bold_prefix=f"* Function: `{name}` — ")
        
    add_styled_heading("3.2.1 Function Source Code Examples", level=2)
    add_p("Below is the source code for the key calculation and security functions defined in our Oracle database:", bold_prefix="Implementation: ")
    
    add_p("fn_GetOutstandingFee Function:", bold_prefix="1. ")
    add_code_block("""CREATE OR REPLACE FUNCTION fn_GetOutstandingFee(p_student_id NUMBER) RETURN NUMBER IS
  l_total_due NUMBER;
  l_total_paid NUMBER;
BEGIN
  SELECT NVL(SUM(fs.amount),0)
    INTO l_total_due
    FROM fee_structures fs
    JOIN students s ON s.program_id = fs.program_id
   WHERE s.student_id = p_student_id;

  SELECT NVL(SUM(fp.amount),0)
    INTO l_total_paid
    FROM fee_payments fp
   WHERE fp.student_id = p_student_id;

  RETURN GREATEST(l_total_due - l_total_paid, 0);
END;""")

    add_p("fn_CalculateCGPA Function:", bold_prefix="2. ")
    add_code_block("""CREATE OR REPLACE FUNCTION fn_CalculateCGPA(p_student_id NUMBER) RETURN NUMBER IS
  l_total_grade_credits NUMBER;
  l_total_credits NUMBER;
BEGIN
  SELECT NVL(SUM(g.grade_points * c.credit_hours), 0), NVL(SUM(c.credit_hours), 0)
    INTO l_total_grade_credits, l_total_credits
    FROM grades g
    JOIN enrollments e ON g.enrollment_id = e.enrollment_id
    JOIN sections sec ON e.section_id = sec.section_id
    JOIN courses c ON sec.course_id = c.course_id
   WHERE e.student_id = p_student_id;
  IF l_total_credits = 0 THEN
    RETURN 0;
  END IF;
  RETURN ROUND(l_total_grade_credits / l_total_credits, 2);
END;""")
        
    # Triggers
    add_styled_heading("3.3 Triggers", level=2)
    add_p("Triggers automatically fire business logic when DML operations (INSERT, UPDATE, DELETE) occur. The schema defines 6 triggers:", bold_prefix="Purpose: ")
    triggers_list = [
        ("trg_AfterFeePayment", "An AFTER INSERT trigger logging fee payments into the system audit_log."),
        ("trg_AuditStudentChanges", "An AFTER INSERT/UPDATE/DELETE trigger recording student record changes in the audit_log."),
        ("trg_AuditGradeChanges", "An AFTER INSERT/UPDATE/DELETE trigger tracking student grade adjustments in the audit_log."),
        ("trg_AuditFeePaymentChanges", "An AFTER INSERT/UPDATE/DELETE trigger monitoring fee payment modifications in the audit_log."),
        ("trg_AfterLibraryReturn", "An AFTER UPDATE trigger incrementing a book's available copies when marked returned."),
        ("trg_InsteadOfStudentUpdate", "An INSTEAD OF trigger on vw_StudentDashboard routing updates to the physical students table.")
    ]
    for name, desc in triggers_list:
        add_p(desc, bold_prefix=f"* Trigger: `{name}` — ")
        
    add_styled_heading("3.3.1 Trigger Source Code Examples", level=2)
    add_p("Below is the source code for our audit and view routing triggers, showing trigger bindings and contexts:", bold_prefix="Implementation: ")
    
    add_p("trg_AuditStudentChanges AFTER Trigger:", bold_prefix="1. ")
    add_code_block("""CREATE OR REPLACE TRIGGER trg_AuditStudentChanges
AFTER INSERT OR UPDATE OR DELETE ON students
FOR EACH ROW
DECLARE
  v_op VARCHAR2(10);
BEGIN
  IF INSERTING THEN
    v_op := 'INSERT';
  ELSIF UPDATING THEN
    v_op := 'UPDATE';
  ELSE
    v_op := 'DELETE';
  END IF;
  INSERT INTO audit_log(table_name, operation, record_id, old_values, new_values, changed_by)
  VALUES (
    'STUDENTS',
    v_op,
    NVL(TO_CHAR(:OLD.student_id), TO_CHAR(:NEW.student_id)),
    CASE WHEN :OLD.student_id IS NOT NULL THEN 'name=' || :OLD.first_name || ' ' || :OLD.last_name ELSE NULL END,
    CASE WHEN :NEW.student_id IS NOT NULL THEN 'name=' || :NEW.first_name || ' ' || :NEW.last_name ELSE NULL END,
    SYS_CONTEXT('USERENV','SESSION_USER')
  );
END;""")

    add_p("trg_InsteadOfStudentUpdate INSTEAD OF Trigger:", bold_prefix="2. ")
    add_code_block("""CREATE OR REPLACE TRIGGER trg_InsteadOfStudentUpdate
INSTEAD OF UPDATE ON vw_StudentDashboard
FOR EACH ROW
BEGIN
  UPDATE students
     SET first_name = SUBSTR(:NEW.student_name, 1, INSTR(:NEW.student_name, ' ') - 1),
         last_name  = SUBSTR(:NEW.student_name, INSTR(:NEW.student_name, ' ') + 1)
   WHERE student_id = :OLD.student_id;
END;""")
        
    # Views
    add_styled_heading("3.4 Database Views", level=2)
    add_p("Views simplify application-tier queries and provide security. The system defines 13 views:", bold_prefix="Purpose: ")
    views_list = [
        ("vw_StudentDashboard", "Displays CGPA, attendance percentage, and outstanding fees for student dashboards."),
        ("vw_AvailableCourses", "Lists sections and courses available for student registration, filtering out already enrolled classes."),
        ("vw_FacultySections", "Lists course sections and schedules assigned to faculty members."),
        ("vw_SectionStudents", "Associates student profiles with sections for faculty roster checks."),
        ("vw_FacultyWorkload", "Aggregates teaching workloads per faculty member, ranking sections by class size."),
        ("vw_FacultyCourseLoad", "Displays sections, courses, and schedules assigned to teachers."),
        ("vw_DepartmentEnrollmentSummary", "Summarizes total enrollments and unique student registrations in each department."),
        ("vw_FeeDefaulters", "Lists students who have outstanding fee balances greater than zero."),
        ("vw_AttendanceShortfall", "Lists students whose class attendance percentage has fallen below 75%."),
        ("vw_LibraryOverdue", "Displays checkout details for library books that are past their due dates."),
        ("vw_ExamTimetable", "Lists scheduled exams, dates, times, rooms, and courses."),
        ("vw_ResultCard", "Combines grades, letters, and credits per section for student results checklists."),
        ("vw_SemesterAttendanceMatrix", "Pivots student attendance percentages per semester into columns.")
    ]
    for name, desc in views_list:
        add_p(desc, bold_prefix=f"* View: `{name}` — ")
        
    add_styled_heading("3.4.1 View Source Code Examples", level=2)
    add_p("Below is the source code for our dashboard and pivoted reports views:", bold_prefix="Implementation: ")
    
    add_p("vw_StudentDashboard Definition:", bold_prefix="1. ")
    add_code_block("""CREATE OR REPLACE VIEW vw_StudentDashboard AS
SELECT s.student_id,
       s.first_name || ' ' || s.last_name AS student_name,
       fn_CalculateCGPA(s.student_id) AS cgpa,
       fn_GetAttendancePercentage(s.student_id) AS attendance_percent,
       fn_GetOutstandingFee(s.student_id) AS outstanding_fees
  FROM students s;""")

    add_p("vw_SemesterAttendanceMatrix (Pivoted View) Definition:", bold_prefix="2. ")
    add_code_block("""CREATE OR REPLACE VIEW vw_SemesterAttendanceMatrix AS
SELECT *
  FROM (
    SELECT s.student_id,
           sec.semester,
           fn_GetAttendancePercentage(s.student_id) attendance_percent
      FROM students s
      JOIN enrollments e ON s.student_id = e.student_id
      JOIN sections sec ON e.section_id = sec.section_id
  )
  PIVOT (
    MAX(attendance_percent) FOR semester IN ('Spring' AS spring, 'Fall' AS fall, 'Summer' AS summer)
  );""")

    doc.add_page_break()

    # PAGE 19-21: DATABASE ARCHITECTURE & ERD
    print("Writing Chapter 4: Database Architecture & ERD...")
    add_styled_heading("4. Database Architecture, ERD and Technologies", level=1)
    add_p("The database architecture of HiSUP represents a relational network designed to ensure data consistency, constraint validation, and row-level access security. Tables are linked to each other using explicit primary and foreign key associations. For instance, `departments` links to `programs` via `dept_id`. The `programs` table links to `students` via `program_id` to govern tuition billing structures, and `user_accounts` maintains a 1-to-1 link with both `students` and `faculty` tables via their shared keys to secure login profiles.", bold_prefix="Database Architecture: ")
    add_p("In Oracle, the ERD generation was managed automatically. The Asgard developer database tool was utilized to reverse-engineer the live relational schema. By connecting the Asgard developer tool to the Oracle XE schema, it automatically generated the entity-relationship model, tracing the cardinalities, indexes, and primary/foreign key connections directly from the database catalog. This automatically generated model is depicted in the diagram below.", bold_prefix="ERD Generation via Asgard developer: ")
    add_screenshot_placeholder("Automatically Generated ERD in Asgard developer")
    
    add_styled_heading("4.1 Technology Stack & Connection Management", level=2)
    add_p("The project is developed using the following technologies:", bold_prefix="Technologies: ")
    add_bullet("A responsive UI client built with component structures and styled using Bootstrap 5.", bold_prefix="Frontend: ")
    add_bullet("An asynchronous server routing requests, handling session tokens, and wrapping endpoints.", bold_prefix="Backend: ")
    add_bullet("An enterprise relational engine executing stored procedures and managing access policies.", bold_prefix="Database: ")
    add_bullet("A binary driver allowing Express to connect to Oracle via high-speed socket streams.", bold_prefix="Database Client: ")
    
    add_p("The database connection string is configured securely on the backend. In development, a local connection string targets Oracle XE. In production, a secure environment variable on the hosting platform provides the cloud connection string, ensuring that credentials never appear in plain text in the GitHub repository. Connection pools are initialized at server startup to ensure efficient allocation of database sessions among concurrent client requests.", bold_prefix="Connection String Management: ")

    doc.add_page_break()

    # PAGE 22-23: VERSION CONTROL & GITHUB
    print("Writing Chapter 5: Version Control & GitHub...")
    add_styled_heading("5. Version Control & GitHub Repository", level=1)
    add_p("Software development for HiSUP was managed under version control using Git. A public repository was maintained throughout the development timeline to track changes, maintain code backups, and support cooperative development workflows.", bold_prefix="Repository Overview: ")
    add_p("To show steady, incremental progress and maintain clean repository history, multiple commits were made on GitHub after completing each meaningful unit of work (e.g., adding a table, creating a stored procedure, routing an Express controller, or styling a frontend view). Commit messages follow a structured prefix notation (such as `[DB]`, `[Proc]`, `[Web]`, or `[Setup]`) to make the version history readable and easy to audit.", bold_prefix="Incremental Commits: ")
    add_screenshot_placeholder("Git Commit History List on GitHub")

    doc.add_page_break()

    # PAGE 24-25: CHALLENGES & REFLECTIONS
    print("Writing Chapter 6: Challenges & Reflections...")
    add_styled_heading("6. Challenges, Lessons Learned and Reflection", level=1)
    
    add_styled_heading("6.1 Technical Challenges", level=2)
    add_p("Establishing a secure identity context inside a stateless REST API (Node/Express) was a challenge. Since connection pools reuse database sessions, database security settings can leak between requests if not managed properly. We resolved this by setting the database `CLIENT_IDENTIFIER` dynamically at the start of each request, and clearing it in a final block when the connection is released.", bold_prefix="Challenge 1: Connection Pooling & VPD Identity Leakage — ")
    add_p("Implementing column-level encryption in a way that allows the application to query data efficiently was another challenge. If encryption is performed using random keys at the application tier, database indexes cannot be used for searches. We resolved this by using a deterministic encryption function (`fn_encrypt_value`) in PL/SQL. This allows the database to perform equality matches on encrypted columns directly (e.g. looking up a student by an encrypted CNIC) without decrypting every row.", bold_prefix="Challenge 2: Column Encryption & Search Performance — ")
    add_p("Writing safe dynamic SQL that is immune to SQL injection while keeping queries readable was a challenge. We avoided string concatenation by using Oracle's `SYS_REFCURSOR` and executing parameterised dynamic statements using the `USING` clause. This allows the query optimizer to cache query plans and reuse them, preventing SQL injection.", bold_prefix="Challenge 3: Safe Parameterised Dynamic SQL — ")
    add_p("Handling deadlocks during peak registration hours was a challenge. In Oracle, deadlocks throw `ORA-00060` and roll back the blocked statement. We resolved this by implementing an exponential backoff retry loop in the Express backend, catching the deadlock error and retrying the operation up to three times before returning an error to the user.", bold_prefix="Challenge 4: Deadlocks in Concurrent Registrations — ")
    
    add_styled_heading("6.2 Reflection", level=2)
    add_p("This project bridged the gap between theoretical database concepts taught in lectures and practical, real-world deployment challenges. In lectures, normalization, transactions, and row-level security are discussed in isolation. Building a complete, working university portal forced us to understand how these concepts interact. For instance, we saw how normalization affects query complexity, requiring joins that must be optimized with indexes, and how row-level security policies must be integrated with the application server's authentication context.", bold_prefix="Reflection: ")
    add_p("Working with Oracle PL/SQL also taught us the importance of keeping business logic close to the database. Moving logic (like GPA calculations and grade updates) from the application server to stored procedures and triggers reduced network latency, simplified backend routing, and ensured data integrity across concurrent sessions. This project provided invaluable hands-on experience in building secure, scalable database applications.", bold_prefix="Conclusion: ")

    # Save output to docs/project_report.docx
    os.makedirs("docs", exist_ok=True)
    out_path = "docs/project_report.docx"
    print(f"Saving document to {out_path}...")
    try:
        doc.save(out_path)
        print("Document saved successfully!")
    except PermissionError:
        alt_path = "docs/project_report_v2.docx"
        print(f"Permission Denied: '{out_path}' is open in Word. Saving as '{alt_path}' instead...")
        doc.save(alt_path)
        print("Document saved successfully to fallback location!")

if __name__ == "__main__":
    create_report()
